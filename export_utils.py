from fpdf import FPDF
from docx import Document

def export_transcript_pdf(transcript, filename="transcript.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for entry in transcript:
        pdf.multi_cell(0, 10, f"[{entry['start']:.2f}-{entry['end']:.2f}] {entry['text']}")
    pdf.output(filename)

def export_transcript_docx(transcript, filename="transcript.docx"):
    doc = Document()
    doc.add_heading("YouTube Transcript", 0)
    for entry in transcript:
        doc.add_paragraph(f"[{entry['start']:.2f}-{entry['end']:.2f}] {entry['text']}")
    doc.save(filename)
